import pandas as pd
import numpy as np
import os
import matplotlib.pyplot as plt
from scipy.stats import shapiro, pearsonr
from statsmodels.formula.api import mixedlm
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches

# Define file paths
imu_data_raw = r"C:\Users\Palan\OneDrive\Documents\School\McMaster\1.Master's Thesis\Thesis_Data_Analysis\I.M.U data\StepSessionSummaryExport"
force_plate_path = r"C:\Users\Palan\OneDrive\Documents\School\McMaster\1.Master's Thesis\Thesis_Data_Analysis\Forceplate Data\forcedecks-test-export-full_v1.xlsx"
srpe_path = r"C:\Users\Palan\OneDrive\Documents\School\McMaster\1.Master's Thesis\Thesis_Data_Analysis\SRPE\McMaster MBB 2023 - Sessional RPE (Study Fall-2023) (Responses) - Form Responses 1.xlsx"

# Load IMU data
imu_data = pd.DataFrame()
for subfolder in os.listdir(imu_data_raw):
    subfolder_path = os.path.join(imu_data_raw, subfolder)
    if os.path.isdir(subfolder_path):
        for date_folder in os.listdir(subfolder_path):
            date_folder_path = os.path.join(subfolder_path, date_folder)
            if os.path.isdir(date_folder_path):
                for root, dirs, files in os.walk(date_folder_path):
                    for file in files:
                        if file.endswith(".csv"):
                            imu_sheet_path = os.path.join(root, file)
                            temp_imu_table = pd.read_csv(imu_sheet_path)
                            imu_data = pd.concat([imu_data, temp_imu_table], ignore_index=True)

# Load Force Plate and SRPE data
force_plate = pd.read_excel(force_plate_path)
srpe = pd.read_excel(srpe_path)

# Combine "First Name" and "Last Name" in IMU data
imu_data['Name'] = imu_data['First Name'] + ' ' + imu_data['Last Name']
imu_data.rename(columns={'Date (YYYY-MM-DD)': 'Date'}, inplace=True)

# Replace specific names
name_replacements = {'Kazim Raza1': 'Kazim Raza', 'Brendan Amoyaw1': 'Brendan Amoyaw', 'Moody Mohamud': 'Moody Muhammed', 'Thomas Matsell': 'Thomas Mattsel'}
imu_data['Name'].replace(name_replacements, inplace=True)
srpe['Name'].replace(name_replacements, inplace=True)

# Select relevant columns
imu_data = imu_data[['Name', 'Date', 'Footnote', 'Impact Load Total (L+R)', 'Average Intensity (L and R)']]
force_plate = force_plate[['Name', 'Date', 'Tags', 'RSI-modified (Imp-Mom) [m/s]', 'Eccentric Duration [ms]', 'Jump Height (Imp-Mom) in Inches [in]', 'Eccentric Braking RFD / BM [N/s/kg]', 'Force at Zero Velocity / BM [N/kg]']]
srpe = srpe[['Name', 'Date', 'SRPE']]

# Standardize date formats
imu_data['Date'] = pd.to_datetime(imu_data['Date']).dt.strftime('%Y-%m-%d')
force_plate['Date'] = pd.to_datetime(force_plate['Date'], errors='coerce').dt.strftime('%Y-%m-%d')
srpe['Date'] = pd.to_datetime(srpe['Date'], errors='coerce').dt.strftime('%Y-%m-%d')

# Merge datasets
complete_data = imu_data.merge(force_plate, on=['Name', 'Date'], how='outer').merge(srpe, on=['Name', 'Date'], how='outer')
all_footnote_data = complete_data[complete_data['Footnote'] == 'All']

# Filter and rename columns
pre_data = all_footnote_data[all_footnote_data['Tags'] == 'Pre-Practice']
post_data = all_footnote_data[all_footnote_data['Tags'] == 'Post-Practice']
metric_columns = [col for col in post_data.columns if col not in ['Name', 'Date', 'Tags']]
post_data = post_data.rename(columns={metric: f'{metric}_Post_Practice' for metric in metric_columns})
columns_to_remove = ['Footnote_Post_Practice', 'Impact Load Total per Minute (L and R)_Post_Practice']
post_data = post_data.drop(columns=[col for col in columns_to_remove if col in post_data.columns], errors='ignore')
pre_post_data = pre_data.merge(post_data, on=['Name', 'Date'], suffixes=('_Pre', '_Post'), how='inner')

# Map names to positions
positions = {'Brendan Amoyaw': 'Forward', 'Cashius McNeilly': 'Guard', 'Daniel Graham': 'Guard', 'Elijah Bethune': 'Guard', 'Jeremiah Francis': 'Guard', 'Kazim Raza': 'Guard', 'Matthew Groe': 'Guard', 'Mike Demagus': 'Guard', 'Moody Muhammed': 'Forward', 'Parker Davis': 'Guard', 'Riaz Saliu': 'Forward', 'Sebastian Di Manno': 'Guard', 'Stevan Japundzic': 'Forward', 'Thomas Mattsel': 'Forward'}
pre_post_data['Position'] = pre_post_data['Name'].map(positions)
unique_names = pre_post_data['Name'].unique()
player_ids = {name: idx for idx, name in enumerate(unique_names, 1)}
pre_post_data['PlayerID'] = pre_post_data['Name'].map(player_ids)
final_data_set = pre_post_data.drop(columns=['Name'])

# Rename columns to avoid spaces and special characters
final_data_set.rename(columns={
    'Impact Load Total (L+R)': 'Impact_Load_Total_LR',
    'Average Intensity (L and R)': 'Average_Intensity_LR',
    'RSI-modified (Imp-Mom) [m/s]': 'RSI_modified_Imp_Mom_m_s',
    'Eccentric Duration [ms]': 'Eccentric_Duration_ms',
    'Jump Height (Imp-Mom) in Inches [in]': 'Jump_Height_Imp_Mom_in_Inches_in',
    'Eccentric Braking RFD / BM [N/s/kg]': 'Eccentric_Braking_RFD_BM_N_s_kg',
    'Force at Zero Velocity / BM [N/kg]': 'Force_at_Zero_Velocity_BM_N_kg',
    'Impact Load Total (L+R)_Post_Practice': 'Impact_Load_Total_LR_Post_Practice',
    'Average Intensity (L and R)_Post_Practice': 'Average_Intensity_LR_Post_Practice',
    'RSI-modified (Imp-Mom) [m/s]_Post_Practice': 'RSI_modified_Imp_Mom_m_s_Post_Practice',
    'Eccentric Duration [ms]_Post_Practice': 'Eccentric_Duration_ms_Post_Practice',
    'Jump Height (Imp-Mom) in Inches [in]_Post_Practice': 'Jump_Height_Imp_Mom_in_Inches_in_Post_Practice',
    'Eccentric Braking RFD / BM [N/s/kg]_Post_Practice': 'Eccentric_Braking_RFD_BM_N_s_kg_Post_Practice',
    'Force at Zero Velocity / BM [N/kg]_Post_Practice': 'Force_at_Zero_Velocity_BM_N_kg_Post_Practice',
    'SRPE_Post-Practice': 'SRPE_Post_Practice'
}, inplace=True)

# Variables for the mixed model
variables_for_model = [
    'Impact_Load_Total_LR', 
    'Average_Intensity_LR', 
    'RSI_modified_Imp_Mom_m_s', 
    'Eccentric_Duration_ms', 
    'Jump_Height_Imp_Mom_in_Inches_in', 
    'Eccentric_Braking_RFD_BM_N_s_kg', 
    'Force_at_Zero_Velocity_BM_N_kg', 
    'SRPE', 
    'Impact_Load_Total_LR_Post_Practice', 
    'Average_Intensity_LR_Post_Practice', 
    'RSI_modified_Imp_Mom_m_s_Post_Practice', 
    'Eccentric_Duration_ms_Post_Practice', 
    'Jump_Height_Imp_Mom_in_Inches_in_Post_Practice', 
    'Eccentric_Braking_RFD_BM_N_s_kg_Post_Practice', 
    'Force_at_Zero_Velocity_BM_N_kg_Post_Practice', 
    'SRPE_Post_Practice'
]
final_data_set = final_data_set[['PlayerID', 'Date', 'Footnote', 'Position'] + variables_for_model]

# Melt the dataframe
linear_mixed_model = pd.melt(final_data_set, id_vars=['PlayerID', 'Date', 'Footnote', 'Position', 'Impact_Load_Total_LR', 'Average_Intensity_LR', 'SRPE'], value_vars=variables_for_model, var_name='Measurement_Type', value_name='Value')
linear_mixed_model['Time'] = linear_mixed_model['Measurement_Type'].apply(lambda x: 'Post' if 'Post' in x else 'Pre')
linear_mixed_model['Measurement_Type'] = linear_mixed_model['Measurement_Type'].str.replace('_Post_Practice', '')
linear_mixed_model['Measurement_Type'] = linear_mixed_model['Measurement_Type'].str.replace('-Practice', '')  # Remove '-Practice'
linear_mixed_model = linear_mixed_model[~linear_mixed_model['Measurement_Type'].isin(['Impact_Load_Total_LR', 'Average_Intensity_LR'])]
linear_mixed_model = linear_mixed_model.drop_duplicates()

# Outlier removal using IQR
def remove_outliers_iqr(df, value_vars):
    for var in value_vars:
        for time in ['Pre', 'Post']:
            subset = df[(df['Measurement_Type'] == var) & (df['Time'] == time)]
            Q1 = subset['Value'].quantile(0.25)
            Q3 = subset['Value'].quantile(0.75)
            IQR = Q3 - Q1
            lower_bound = Q1 - 1.5 * IQR
            upper_bound = Q3 + 1.5 * IQR
            df = df[~((df['Measurement_Type'] == var) & (df['Time'] == time) & ((df['Value'] < lower_bound) | (df['Value'] > upper_bound)))]
    return df

# Remove outliers
value_vars = linear_mixed_model['Measurement_Type'].unique()
linear_mixed_model = remove_outliers_iqr(linear_mixed_model, value_vars)

# Shapiro-Wilk test for normality
shapiro_results = []
for measurement in linear_mixed_model['Measurement_Type'].unique():
    subset = linear_mixed_model[linear_mixed_model['Measurement_Type'] == measurement]
    if len(subset) > 0:
        stat, p = shapiro(subset['Value'])
        shapiro_results.append({'Measurement_Type': measurement, 'Statistic': stat, 'p-value': p, 'Significant': 'Yes' if p <= 0.05 else 'No'})

shapiro_results_df = pd.DataFrame(shapiro_results)

# Calculate descriptive statistics
descriptive_stats = linear_mixed_model.groupby(['Measurement_Type', 'Time'])['Value'].describe().unstack()

# Ensure matching lengths for correlation analysis
def get_matching_data(x, y):
    mask = ~x.isna() & ~y.isna()
    return x[mask], y[mask]

# Create a Word document to save results
doc = Document()
doc.add_heading('Mixed Model Analysis Results', 0)

# Final analysis function
def final_analysis(data, by='Overall'):
    measurement_types = data['Measurement_Type'].unique()
    
    doc.add_heading(f'Processing {by}:\n', level=1)
    
    for var in measurement_types:
        data_subset = data[data['Measurement_Type'] == var]
        try:
            model = mixedlm("Value ~ Impact_Load_Total_LR + Average_Intensity_LR + Time", data_subset, groups=data_subset["PlayerID"])
            result = model.fit()
            doc.add_heading(f'Results for {var} ({by}):', level=2)
            doc.add_paragraph(result.summary().as_text())
            
            if 'Impact_Load_Total_LR' in result.pvalues.index and result.pvalues['Impact_Load_Total_LR'] < 0.05:
                doc.add_paragraph(f"    -> The effect of Impact Load Total (L+R) on {var} ({by}) is significant.")
            else:
                doc.add_paragraph(f"    -> The effect of Impact Load Total (L+R) on {var} ({by}) is not significant.")
            if 'Average_Intensity_LR' in result.pvalues.index and result.pvalues['Average_Intensity_LR'] < 0.05:
                doc.add_paragraph(f"    -> The effect of Average Intensity (L and R) on {var} ({by}) is significant.")
            else:
                doc.add_paragraph(f"    -> The effect of Average Intensity (L and R) on {var} ({by}) is not significant.")
        except Exception as e:
            doc.add_paragraph(f"Model for {var} ({by}) did not converge: {e}")
        doc.add_paragraph("\n")

# Overall analysis
final_analysis(linear_mixed_model)

# Analysis by position
for position in ['Forward', 'Guard']:
    position_data = linear_mixed_model[linear_mixed_model['Position'] == position]
    final_analysis(position_data, by=position)

# Analysis by date
dates = linear_mixed_model['Date'].unique()
for date in dates:
    date_data = linear_mixed_model[linear_mixed_model['Date'] == date]
    final_analysis(date_data, by=date)

# Save the document
doc.save("Mixed_Model_Analysis_Results.docx")

